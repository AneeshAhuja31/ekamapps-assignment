import os
import asyncio
import sqlite3
from typing import List, Dict, Any, Optional
import json
from datetime import datetime
import hashlib
import logging
from pathlib import Path
import uuid
import pickle

import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
import numpy as np

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.embeddings_model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
        self.db_path = "documents.db"
        self.vector_store_path = "document_vector_store"
        
        self.huggingface_embeddings = None
        self.vector_store = None
        
        try:
            self.huggingface_embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            logger.info("HuggingFace AllMiniLM L6 V2 embeddings initialized successfully")
                
        except Exception as e:
            logger.warning(f"Could not initialize HuggingFace Embeddings: {e}. Vector store will be disabled.")
        
        self._initialize_db()
        self._load_or_create_vector_store()
    
    def _initialize_db(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                file_type TEXT NOT NULL,
                file_size INTEGER,
                content TEXT,
                metadata TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                processed_at TIMESTAMP,
                chunk_count INTEGER DEFAULT 0
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS document_chunks (
                id TEXT PRIMARY KEY,
                document_id TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                content TEXT NOT NULL,
                embedding BLOB,
                metadata TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (document_id) REFERENCES documents (id)
            )
        ''')
        
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_filename ON documents (filename)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_chunk_doc_id ON document_chunks (document_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_chunk_index ON document_chunks (chunk_index)')
        
        conn.commit()
        conn.close()
        
        logger.info("Document database initialized successfully")

    def _load_or_create_vector_store(self):
        try:
            if not self.huggingface_embeddings:
                logger.warning("HuggingFace embeddings not available. Vector store disabled.")
                return
                
            if os.path.exists(self.vector_store_path):
                try:
                    self.vector_store = FAISS.load_local(
                        self.vector_store_path,
                        embeddings=self.huggingface_embeddings,
                        allow_dangerous_deserialization=True
                    )
                    logger.info(f"Loaded existing vector store from {self.vector_store_path}")
                except Exception as load_error:
                    logger.warning(f"Failed to load existing vector store: {load_error}")
                    self._create_empty_vector_store()
            else:
                self._create_empty_vector_store()
                
        except Exception as e:
            logger.error(f"Error initializing vector store: {e}")
            self.vector_store = None
    
    def _create_empty_vector_store(self):
        try:
            placeholder_text = "This is a placeholder document for vector store initialization."
            self.vector_store = FAISS.from_texts(
                [placeholder_text],
                self.huggingface_embeddings,
                metadatas=[{"type": "placeholder", "filename": "placeholder"}]
            )
            self.vector_store.save_local(self.vector_store_path)
            logger.info(f"Created new vector store at {self.vector_store_path}")
        except Exception as e:
            logger.error(f"Failed to create empty vector store: {e}")
            self.vector_store = None

    
    async def process_document(self, file_path: str) -> str:
        try:
            doc_id = str(uuid.uuid4())
            
            file_info = self._get_file_info(file_path)
            
            content = await self._extract_content(file_path, file_info['type'])
            
            if not content or len(content.strip()) < 10:
                raise ValueError(f"Could not extract meaningful content from {file_path}")
            
            chunks = await self._create_intelligent_chunks(content, file_info['type'])
            chunk_embeddings = await self._generate_embeddings(chunks)
            
            await self._store_document(doc_id, file_info, content, chunks, chunk_embeddings)
            
            if self.vector_store and self.huggingface_embeddings:
                await self._store_in_vector_store(doc_id, file_info, chunks)
            
            logger.info(f"Successfully processed document: {file_path} -> {len(chunks)} chunks")
            return doc_id
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        file_size = path.stat().st_size if path.exists() else 0
        
        extension = path.suffix.lower()
        if extension == '.pdf':
            file_type = 'pdf'
        elif extension in ['.docx', '.doc']:
            file_type = 'docx'
        elif extension == '.txt':
            file_type = 'txt'
        elif extension == '.csv':
            file_type = 'csv'
        else:
            file_type = 'unknown'
        
        return {
            'filename': path.name,
            'type': file_type,
            'size': file_size,
            'extension': extension
        }
    
    async def _extract_content(self, file_path: str, file_type: str) -> str:
        """Extract text content from various file types"""
        try:
            if file_type == 'pdf':
                return await self._extract_pdf_content(file_path)
            elif file_type == 'docx':
                return await self._extract_docx_content(file_path)
            elif file_type == 'txt':
                return await self._extract_txt_content(file_path)
            elif file_type == 'csv':
                return await self._extract_csv_content(file_path)
            else:
                return await self._extract_txt_content(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            raise
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
            return content.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF content: {e}")
            return ""
    
    async def _extract_docx_content(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            return "\n".join(content)
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            return ""
    
    async def _extract_txt_content(self, file_path: str) -> str:
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            return ""
        except Exception as e:
            logger.error(f"Error extracting TXT content: {e}")
            return ""
    
    async def _extract_csv_content(self, file_path: str) -> str:
        """Extract text representation from CSV"""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            
            content = f"CSV Data Summary:\n"
            content += f"Columns: {', '.join(df.columns.tolist())}\n"
            content += f"Rows: {len(df)}\n\n"
            
            content += "Sample Data:\n"
            content += df.head(5).to_string(index=False)
            
            return content
        except Exception as e:
            logger.error(f"Error extracting CSV content: {e}")
            return ""
    
    async def _create_intelligent_chunks(self, content: str, doc_type: str) -> List[str]:
        """Create intelligent chunks based on document structure"""
        if not content or len(content.strip()) == 0:
            return []
        
        if doc_type == 'csv':
            chunk_size = 1000  
        elif doc_type in ['pdf', 'docx']:
            chunk_size = 800  
        else:
            chunk_size = 600   
        
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        if not paragraphs:
            paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk += "\n" + paragraph if current_chunk else paragraph
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        if not chunks and content.strip():
            sentences = content.split('. ')
            current_chunk = ""
            for sentence in sentences:
                if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = sentence
                else:
                    current_chunk += ". " + sentence if current_chunk else sentence
            
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
        
        if not chunks:
            for i in range(0, len(content), chunk_size):
                chunk = content[i:i + chunk_size]
                if chunk.strip():
                    chunks.append(chunk.strip())
        
        return chunks
    
    async def _generate_embeddings(self, chunks: List[str]) -> List[np.ndarray]:
        try:
            batch_size = 32
            embeddings = []
            
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                batch_embeddings = self.embeddings_model.encode(batch)
                embeddings.extend(batch_embeddings)
            
            return embeddings
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            raise
    
    async def _store_document(self, doc_id: str, file_info: Dict, content: str, 
                            chunks: List[str], embeddings: List[np.ndarray]):
        """Store document and chunks in database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            metadata = json.dumps({
                'original_filename': file_info['filename'],
                'file_extension': file_info['extension']
            })
            
            cursor.execute('''
                INSERT INTO documents (id, filename, file_type, file_size, content, metadata, processed_at, chunk_count)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                doc_id,
                file_info['filename'],
                file_info['type'],
                file_info['size'],
                content,
                metadata,
                datetime.now().isoformat(),
                len(chunks)
            ))
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_id = f"{doc_id}_chunk_{i}"
                chunk_metadata = json.dumps({
                    'chunk_index': i,
                    'chunk_length': len(chunk)
                })
                
                embedding_bytes = embedding.astype(np.float32).tobytes()
                
                cursor.execute('''
                    INSERT INTO document_chunks (id, document_id, chunk_index, content, embedding, metadata)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (
                    chunk_id,
                    doc_id,
                    i,
                    chunk,
                    embedding_bytes,
                    chunk_metadata
                ))
            
            conn.commit()
            
        except Exception as e:
            conn.rollback()
            logger.error(f"Error storing document: {e}")
            raise
        finally:
            conn.close()
    
    async def _store_in_vector_store(self, doc_id: str, file_info: Dict, chunks: List[str]):
        """Store document chunks in FAISS vector store"""
        try:
            if not self.vector_store or not self.huggingface_embeddings:
                logger.warning("Vector store or embeddings not available")
                return
            
            documents = []
            for i, chunk in enumerate(chunks):
                if len(chunk.strip()) < 5:  #skip very short chunks
                    continue
                    
                doc = Document(
                    page_content=chunk.strip(),
                    metadata={
                        'doc_id': doc_id,
                        'filename': file_info['filename'],
                        'file_type': file_info['type'],
                        'chunk_index': i,
                        'upload_date': datetime.now().isoformat()
                    }
                )
                documents.append(doc)
            
            if documents:
                self.vector_store.add_documents(documents)
                
                self.vector_store.save_local(self.vector_store_path)
                
                logger.info(f"Added {len(documents)} chunks to vector store for document {file_info['filename']}")
            
        except Exception as e:
            logger.error(f"Error storing in vector store: {e}")
    
    async def search_documents(self, query: str, limit: int = 10, similarity_threshold: float = 0.3) -> List[Dict]:
        """Search documents using semantic similarity (prefers FAISS vector store if available)"""
        try:
            if self.vector_store and self.huggingface_embeddings:
                return await self.search_documents_vector(query, limit)
            
            return await self._search_documents_sqlite(query, limit, similarity_threshold)
            
        except Exception as e:
            logger.error(f"Error in document search: {e}")
            return []
    
    async def _search_documents_sqlite(self, query: str, limit: int = 10, similarity_threshold: float = 0.3) -> List[Dict]:
        """Search documents using SQLite embeddings (fallback method)"""
        try:
            query_embedding = self.embeddings_model.encode([query])[0]
            
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT dc.id, dc.document_id, dc.content, dc.embedding, dc.chunk_index,
                       d.filename, d.file_type
                FROM document_chunks dc
                JOIN documents d ON dc.document_id = d.id
            ''')
            
            chunks_data = cursor.fetchall()
            conn.close()
            
            if not chunks_data:
                return []
            
            results = []
            for chunk_data in chunks_data:
                chunk_id, doc_id, content, embedding_bytes, chunk_index, filename, file_type = chunk_data
                
                embedding = np.frombuffer(embedding_bytes, dtype=np.float32)
                
                similarity = np.dot(query_embedding, embedding) / (
                    np.linalg.norm(query_embedding) * np.linalg.norm(embedding)
                )
                
                if similarity > similarity_threshold:
                    results.append({
                        'chunk_id': chunk_id,
                        'document_id': doc_id,
                        'filename': filename,
                        'file_type': file_type,
                        'content': content,
                        'similarity': float(similarity),
                        'chunk_index': chunk_index
                    })
            
            results.sort(key=lambda x: x['similarity'], reverse=True)
            return results[:limit]
            
        except Exception as e:
            logger.error(f"Error searching documents in SQLite: {e}")
            return []
    
    async def search_documents_vector(self, query: str, limit: int = 10) -> List[Dict]:
        """Search documents using FAISS vector store (faster and more accurate)"""
        try:
            if not self.vector_store:
                logger.warning("Vector store not available, falling back to SQLite search")
                return await self._search_documents_sqlite(query, limit)
                
            if not self.huggingface_embeddings:
                logger.warning("HuggingFace embeddings not available, falling back to SQLite search")
                return await self._search_documents_sqlite(query, limit)
            
            logger.debug(f"Performing vector search for query: {query[:50]}...")
            similar_docs = self.vector_store.similarity_search_with_score(query, k=limit)
            logger.debug(f"Vector search returned {len(similar_docs)} results")
            
            results = []
            for doc, score in similar_docs:
                if doc.metadata.get("type") == "placeholder":
                    logger.debug("Skipping placeholder document")
                    continue
                    
                results.append({
                    'document_id': doc.metadata.get('doc_id'),
                    'filename': doc.metadata.get('filename'),
                    'file_type': doc.metadata.get('file_type'),
                    'content': doc.page_content,
                    'similarity': 1.0 - score,  
                    'chunk_index': doc.metadata.get('chunk_index', 0),
                    'upload_date': doc.metadata.get('upload_date')
                })
            
            logger.debug(f"Returning {len(results)} filtered results")
            return results
            
        except Exception as e:
            logger.error(f"Error searching documents with vector store: {e}", exc_info=True)
            return await self._search_documents_sqlite(query, limit)
    

    async def get_total_documents(self) -> int:
        """Get total number of documents"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) FROM documents')
            count = cursor.fetchone()[0]
            conn.close()
            return count
            
        except Exception as e:
            logger.error(f"Error getting total documents: {e}")
            return 0
    
    async def get_total_chunks(self) -> int:
        """Get total number of text chunks"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) FROM document_chunks')
            count = cursor.fetchone()[0]
            conn.close()
            return count
            
        except Exception as e:
            logger.error(f"Error getting total chunks: {e}")
            return 0
    
    async def get_document_info(self, doc_id: str) -> Optional[Dict]:
        """Get information about a specific document"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT id, filename, file_type, file_size, metadata, created_at, processed_at, chunk_count
                FROM documents WHERE id = ?
            ''', (doc_id,))
            
            result = cursor.fetchone()
            conn.close()
            
            if result:
                return {
                    'id': result[0],
                    'filename': result[1],
                    'file_type': result[2],
                    'file_size': result[3],
                    'metadata': json.loads(result[4]) if result[4] else {},
                    'created_at': result[5],
                    'processed_at': result[6],
                    'chunk_count': result[7]
                }
            
            return None
            
        except Exception as e:
            logger.error(f"Error getting document info: {e}")
            return None
    async def get_all_documents(self) -> List[Dict[str, Any]]:
        """Get all processed documents with their content"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT id, filename, content, file_type, processed_at, metadata
                FROM documents
                ORDER BY processed_at DESC
            ''')
            
            documents = []
            for row in cursor.fetchall():
                doc_id, filename, content, file_type, upload_date, metadata_str = row
                
                metadata = {}
                if metadata_str:
                    try:
                        metadata = json.loads(metadata_str)
                    except json.JSONDecodeError:
                        metadata = {}
                
                documents.append({
                    'id': doc_id,
                    'filename': filename,
                    'content': content or "",                     'file_type': file_type,
                    'upload_date': upload_date,
                    'metadata': metadata
                })
            
            conn.close()
            return documents
            
        except Exception as e:
            logger.error(f"Error getting all documents: {e}")
            return []
    
    async def rebuild_vector_store(self):
        try:
            if not self.huggingface_embeddings:
                logger.warning("HuggingFace embeddings not available. Cannot rebuild vector store.")
                return False
            
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT dc.document_id, dc.content, dc.chunk_index, d.filename, d.file_type
                FROM document_chunks dc
                JOIN documents d ON dc.document_id = d.id
                ORDER BY dc.document_id, dc.chunk_index
            ''')
            
            chunks_data = cursor.fetchall()
            conn.close()
            
            if not chunks_data:
                logger.info("No document chunks found. Creating empty vector store.")
                placeholder_text = "This is a placeholder document for vector store initialization."
                self.vector_store = FAISS.from_texts(
                    [placeholder_text],
                    self.huggingface_embeddings,
                    metadatas=[{"type": "placeholder", "filename": "placeholder"}]
                )
            else:
                documents = []
                for doc_id, content, chunk_index, filename, file_type in chunks_data:
                    if len(content.strip()) < 5:  
                        continue
                        
                    doc = Document(
                        page_content=content.strip(),
                        metadata={
                            'doc_id': doc_id,
                            'filename': filename,
                            'file_type': file_type,
                            'chunk_index': chunk_index,
                            'upload_date': datetime.now().isoformat()
                        }
                    )
                    documents.append(doc)
                
                if documents:
                    self.vector_store = FAISS.from_documents(documents, self.huggingface_embeddings)
                    logger.info(f"Rebuilt vector store with {len(documents)} chunks")
                else:
                    logger.warning("No valid chunks found. Creating empty vector store.")
                    placeholder_text = "This is a placeholder document for vector store initialization."
                    self.vector_store = FAISS.from_texts(
                        [placeholder_text],
                        self.huggingface_embeddings,
                        metadatas=[{"type": "placeholder", "filename": "placeholder"}]
                    )
            
            self.vector_store.save_local(self.vector_store_path)
            logger.info(f"Vector store saved to {self.vector_store_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error rebuilding vector store: {e}")
            return False